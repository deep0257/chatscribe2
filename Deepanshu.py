# deepanshu.py

# Load environment variables and setup
from dotenv import load_dotenv
load_dotenv()

# ---- Configuration ----
# Configurations are centralized here
import os
from pydantic_settings import BaseSettings

class Settings(BaseSettings):
    DATABASE_URL: str = os.getenv("DATABASE_URL")
    SECRET_KEY: str = os.getenv("SECRET_KEY")
    ALGORITHM: str = os.getenv("ALGORITHM")
    ACCESS_TOKEN_EXPIRE_MINUTES: int = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES"))
    #OPENAI_API_KEY: str = os.getenv("OPENAI_API_KEY")
    PINECONE_API_KEY: str = os.getenv("PINECONE_API_KEY")
    PINECONE_ENVIRONMENT: str = os.getenv("PINECONE_ENVIRONMENT")
    PINECONE_INDEX_NAME: str = os.getenv("PINECONE_INDEX_NAME")
    DEBUG: bool = os.getenv("DEBUG", "False").lower() == "true"
    HOST: str = os.getenv("HOST")
    PORT: int = int(os.getenv("PORT"))
    UPLOAD_DIR: str = "uploads"
    MAX_FILE_SIZE: int = 10 * 1024 * 1024
    ALLOWED_EXTENSIONS: set = {".pdf", ".docx", ".txt"}

    class Config:
        env_file = ".env"

settings = Settings()

# ---- Models ----
# Database models
from sqlalchemy import Column, Integer, String, DateTime, Text, ForeignKey, Boolean
from sqlalchemy.orm import relationship
from sqlalchemy.sql import func
from sqlalchemy.orm import declarative_base


Base = declarative_base()

class User(Base):
    __tablename__ = "users"

    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, index=True, nullable=False)
    email = Column(String, unique=True, index=True, nullable=False)
    hashed_password = Column(String, nullable=False)
    is_active = Column(Boolean, default=True)
    created_at = Column(DateTime(timezone=True), server_default=func.now())

    documents = relationship("Document", back_populates="owner")
    chat_sessions = relationship("ChatSession", back_populates="user")

class Document(Base):
    __tablename__ = "documents"

    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String, nullable=False)
    original_filename = Column(String, nullable=False)
    file_path = Column(String, nullable=False)
    file_type = Column(String, nullable=False)
    file_size = Column(Integer, nullable=False)
    content = Column(Text)
    uploaded_at = Column(DateTime(timezone=True), server_default=func.now())
    user_id = Column(Integer, ForeignKey("users.id"))

    owner = relationship("User", back_populates="documents")
    chat_sessions = relationship("ChatSession", back_populates="document")

class ChatSession(Base):
    __tablename__ = "chat_sessions"

    id = Column(Integer, primary_key=True, index=True)
    title = Column(String, nullable=False)
    created_at = Column(DateTime(timezone=True), server_default=func.now())
    updated_at = Column(DateTime(timezone=True), server_default=func.now(), onupdate=func.now())
    user_id = Column(Integer, ForeignKey("users.id"))
    document_id = Column(Integer, ForeignKey("documents.id"))

    user = relationship("User", back_populates="chat_sessions")
    document = relationship("Document", back_populates="chat_sessions")
    messages = relationship("ChatMessage", back_populates="session")

class ChatMessage(Base):
    __tablename__ = "chat_messages"

    id = Column(Integer, primary_key=True, index=True)
    content = Column(Text, nullable=False)
    is_user = Column(Boolean, nullable=False)  # True for user messages, False for AI responses
    created_at = Column(DateTime(timezone=True), server_default=func.now())
    session_id = Column(Integer, ForeignKey("chat_sessions.id"))

    session = relationship("ChatSession", back_populates="messages")

# ---- Database ----
# Database setup
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

engine = create_engine(settings.DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

# Dependency

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ---- Schemas ----
# Pydantic Schemas
from datetime import datetime
from typing import List, Optional
from pydantic import BaseModel, EmailStr

class UserBase(BaseModel):
    username: str
    email: EmailStr

class UserCreate(UserBase):
    password: str

class User(UserBase):
    id: int
    is_active: bool
    created_at: datetime

    class Config:
        from_attributes = True

class DocumentBase(BaseModel):
    filename: str
    original_filename: str
    file_type: str

class Document(DocumentBase):
    id: int
    file_size: int
    uploaded_at: datetime
    user_id: int

    class Config:
        orm_mode = True

class ChatMessageBase(BaseModel):
    content: str
    is_user: bool

class ChatMessage(ChatMessageBase):
    id: int
    created_at: datetime
    session_id: int

    class Config:
        orm_mode = True

class ChatSessionBase(BaseModel):
    title: str

class ChatSession(ChatSessionBase):
    id: int
    created_at: datetime
    updated_at: datetime
    user_id: int
    document_id: int
    messages: List[ChatMessage] = []

    class Config:
        orm_mode = True

class Token(BaseModel):
    access_token: str
    token_type: str

class ChatRequest(BaseModel):
    message: str
    session_id: int

class ChatResponse(BaseModel):
    response: str
    session_id: int

# ---- Security ----
# Security utilities
from datetime import timedelta
from jose import JWTError, jwt
from passlib.context import CryptContext

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, settings.SECRET_KEY, algorithm=settings.ALGORITHM)
    return encoded_jwt

def verify_token(token: str) -> Optional[str]:
    try:
        payload = jwt.decode(token, settings.SECRET_KEY, algorithms=[settings.ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            return None
        return username
    except JWTError:
        return None

# ---- CRUD ----
# CRUD operations
from sqlalchemy.orm import Session

def create_user(db: Session, user: UserCreate):
    hashed_password = get_password_hash(user.password)
    db_user = User(username=user.username, email=user.email, hashed_password=hashed_password)
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    return db_user

def authenticate_user(db: Session, username: str, password: str):
    user = db.query(User).filter(User.username == username).first()
    if not user:
        return None
    if not verify_password(password, user.hashed_password):
        return None
    return user

# ---- File Processor ----
# File processing utilities
from pathlib import Path
import uuid
import PyPDF2
import docx
from typing import Tuple, Optional

class FileProcessor:
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)

    def is_allowed_file(self, filename: str) -> bool:
        return Path(filename).suffix.lower() in settings.ALLOWED_EXTENSIONS

    def generate_unique_filename(self, original_filename: str) -> str:
        ext = Path(original_filename).suffix
        unique_name = f"{uuid.uuid4()}{ext}"
        return unique_name

    def save_file(self, file_content: bytes, filename: str) -> str:
        file_path = self.upload_dir / filename
        with open(file_path, "wb") as f:
            f.write(file_content)
        return str(file_path)

    def extract_text_from_pdf(self, file_path: str) -> Optional[str]:
        try:
            text = ""
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None

    def extract_text_from_docx(self, file_path: str) -> Optional[str]:
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return None

    def extract_text_from_txt(self, file_path: str) -> Optional[str]:
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read().strip()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="latin-1") as file:
                    return file.read().strip()
            except Exception as e:
                print(f"Error reading TXT file: {e}")
                return None
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            return None

    def extract_text(self, file_path: str, file_type: str) -> Optional[str]:
        if file_type.lower() == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_type.lower() == ".docx":
            return self.extract_text_from_docx(file_path)
        elif file_type.lower() == ".txt":
            return self.extract_text_from_txt(file_path)
        else:
            return None

    def process_uploaded_file(self, file_content: bytes, original_filename: str) -> Tuple[Optional[str], Optional[str], Optional[str]]:
        try:
            if not self.is_allowed_file(original_filename):
                return None, None, None

            unique_filename = self.generate_unique_filename(original_filename)

            file_path = self.save_file(file_content, unique_filename)

            file_type = Path(original_filename).suffix.lower()
            extracted_text = self.extract_text(file_path, file_type)

            return file_path, unique_filename, extracted_text

        except Exception as e:
            print(f"Error processing uploaded file: {e}")
            return None, None, None

    def delete_file(self, file_path: str) -> bool:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except Exception as e:
            print(f"Error deleting file: {e}")
            return False

file_processor = FileProcessor()

# ---- AI Services ----
# AI Processing utilities
import openai

class AIService:
    #def __init__(self):
       # openai.api_key = settings.OPENAI_API_KEY

    def summarize_document(self, content: str) -> str:
        try:
            max_tokens = 3000
            if len(content) > max_tokens:
                content = content[:max_tokens]

            prompt = f"""
            Please provide a comprehensive summary of the following document:

            {content}

            Summary:
            """

            response = openai.Completion.create(
                engine="text-davinci-003",
                prompt=prompt,
                max_tokens=500,
                temperature=0.3
            )

            return response.choices[0].text.strip()

        except Exception as e:
            print(f"Error in summarization: {e}")
            return "Error generating summary."

ai_service = AIService()

# ---- FastAPI Setup ----
# FastAPI Application
from fastapi import FastAPI
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.sessions import SessionMiddleware

from fastapi.responses import RedirectResponse
from fastapi.templating import Jinja2Templates
from fastapi import Depends, HTTPException, status, UploadFile, File, APIRouter, Request, Response

app = FastAPI(title="AI-Powered Document Chatbot")

templates = Jinja2Templates(directory="app/templates")

# Middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Change to specific domains in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.add_middleware(SessionMiddleware, secret_key=settings.SECRET_KEY)

# Static files
app.mount("/static", StaticFiles(directory="app/static"), name="static")


# ---- Endpoints ----
# Auth dependent utilities
def require_auth(request: Request, db: Session = Depends(get_db), access_token: Optional[str] = None) -> User:
    if not access_token:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Authentication required"
        )

    username = verify_token(access_token)
    if username is None:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or expired token"
        )

    user = db.query(User).filter(User.username == username).first()
    if user is None:
        raise HTTPException(status_code=404, detail="User not found")

    return user

# API Routes
api_router = APIRouter()

@api_router.post("/auth/signup", response_model=Token)
async def signup(user: UserCreate, db: Session = Depends(get_db)):
    if db.query(User).filter(User.username == user.username).first():
        raise HTTPException(status_code=400, detail="Username already registered")

    if db.query(User).filter(User.email == user.email).first():
        raise HTTPException(status_code=400, detail="Email already registered")

    db_user = create_user(db, user)

    access_token_expires = timedelta(minutes=settings.ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(data={"sub": db_user.username}, expires_delta=access_token_expires)

    return {"access_token": access_token, "token_type": "bearer"}

@api_router.post("/auth/login", response_model=Token)
async def login(user_credentials: UserCreate, db: Session = Depends(get_db)):
    user = authenticate_user(db, user_credentials.username, user_credentials.password)
    if not user:
        raise HTTPException(status_code=401, detail="Incorrect username or password")

    access_token_expires = timedelta(minutes=settings.ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(data={"sub": user.username}, expires_delta=access_token_expires)

    return {"access_token": access_token, "token_type": "bearer"}

@api_router.post("/documents/upload")
async def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    file_content = await file.read()
    if len(file_content) > settings.MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large. Maximum size: {settings.MAX_FILE_SIZE / (1024*1024):.1f} MB")

    file_path, unique_filename, extracted_text = file_processor.process_uploaded_file(file_content, file.filename)

    if not file_path or not extracted_text:
        raise HTTPException(status_code=500, detail="Failed to process file")

    document = Document(
        filename=unique_filename,
        original_filename=file.filename,
        file_path=file_path,
        file_type=file.filename.split('.')[-1].lower(),
        file_size=len(file_content),
        content=extracted_text,
        user_id=current_user.id
    )

    db.add(document)
    db.commit()
    db.refresh(document)

    return document

@api_router.get("/documents/", response_model=List[Document])
async def get_documents(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    return db.query(Document).filter(Document.user_id == current_user.id).all()

@api_router.get("/documents/{document_id}", response_model=Document)
async def get_document_detail(document_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    document = db.query(Document).filter(Document.id == document_id, Document.user_id == current_user.id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return document

@api_router.post("/documents/{document_id}/summarize")
async def summarize_document(document_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    document = db.query(Document).filter(Document.id == document_id, Document.user_id == current_user.id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    if not document.content:
        raise HTTPException(status_code=400, detail="Document content not available")

    summary = ai_service.summarize_document(document.content)
    return {"summary": summary}

@api_router.post("/chat/start")
async def start_new_chat(chat_request: ChatRequest, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    document_id = chat_request.session_id
    document = db.query(Document).filter(Document.id == document_id, Document.user_id == current_user.id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    response = ai_service.summarize_document(document.content)
    return {"response": response, "session_id": document_id}

app.include_router(api_router, prefix="/api")

# Web Endpoints
@api_router.get("/", response_class=RedirectResponse)
async def home(request: Request, user: Optional[User] = Depends(require_auth)):
    if user:
        return RedirectResponse(url="/dashboard")
    return templates.TemplateResponse("home.html", {"request": request})

# ---- Running the app ----
# Run using Uvicorn
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

